import gspread
from datetime import datetime
from docx import Document
from verificarLaudos import laudo_ja_gerado
from oauth2client.service_account import ServiceAccountCredentials
import os
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import datetime
from pydrive.auth import GoogleAuth
from pydrive.drive import GoogleDrive
import unicodedata
from unidecode import unidecode
from descriptions import describe_aguia, describe_gato, describe_tubarao, describe_lobo
from logger import gerar_log
import json

# Carregar configurações do arquivo JSON
with open('config.json', 'r') as config_file:
    config = json.load(config_file)

# Acessar o diretório dos laudos e outros valores
diretorio_laudos = config["diretorio_laudos"]


# Defina a variável gc globalmente
gc = gspread.service_account(filename='clean-beaker-414415-8887a7f9338c.json')

def obter_dados_planilha(linha):
    planilha = gc.open_by_url(
        "https://docs.google.com/spreadsheets/d/1U9ZFaFmNf4aJ13gbzxuF8AObk-tDWFOaE9If5iOplNc/edit?resourcekey#gid=307731029"). \
        worksheet("Respostas ao formulário 1")
    dados = planilha.row_values(linha)
    # Remover espaços em branco no início e no final de cada elemento
    dados = [item.strip() if isinstance(item, str) else item for item in dados]
    return dados

def contar_respostas(dados):
    aguia = 0
    gato = 0
    tubarao = 0
    lobo = 0

    # Supondo que os índices das colunas AC até BA sejam 28 até 53
    for resposta in dados[28:54]:  # Colunas AC até BA
        for letra in resposta:
            if letra == 'I':
                aguia += 1
            elif letra == 'C':
                gato += 1
            elif letra == 'A':
                tubarao += 1
            elif letra == 'O':
                lobo += 1

    # Multiplicando cada contagem por 4 e adicionando o símbolo de porcentagem
    aguia_percent = f"{int(aguia) * 4}"
    gato_percent = f"{int(gato) * 4}"
    tubarao_percent = f"{int(tubarao) * 4}"
    lobo_percent = f"{int(lobo) * 4}"

    return aguia_percent, gato_percent, tubarao_percent, lobo_percent

def preencher_informacoes(carimbo, email, usuario, data_nascimento, candidato_vaga ,pretensao_salarial,
endereco, distancia_empresa, meio_transporte, estado_civil, tem_filhos, casa_propria, porque_contratado, porque_identificou,
descricao, tempo_experiencia, porte, nivel, area_atuacao, estudando_atualmente, instituicao_curso_periodo, pacote_office, excel, word, fumante,
dist_empr_temp, filhos_idade, filhos_qtde, seg_ult_empr, temp_ult_empr, mot_saida_ult_empr, seg_penult_empr, temp_penult_empr, mot_saida_penult_empr, telefone, uf_estado, cod_cidade,
aguia, gato, tubarao, lobo, resultado_teste):

    modelo = Document("D:/LaudosTeste/modelo.docx")
    # Obtém a data e hora atual
    now = datetime.now()
    current_time = now.strftime("%d/%m/%Y %H:%M:%S")

    #print(f"{current_time}Nome: {usuario} | Email: {email} | Data Registro: {carimbo}")

    for p in modelo.paragraphs:
        if 'Nome do Candidato:' in p.text:
            p.text = p.text.replace('usuario', f'{usuario}')

        if 'data_nascimento' in p.text:
            idade = calcular_idade(data_nascimento)
            p.text = p.text.replace('data_nascimento', f' {idade} anos')

        if 'Candidato a vaga de:' in p.text:
            p.text = p.text.replace('candidato_vaga', f' {candidato_vaga}')

        if 'Pretensão salarial:' in p.text:
            p.text = p.text.replace('pretensao_salarial', f' {pretensao_salarial}')

        if 'Endereço:' in p.text:
            p.text = p.text.replace('endereco', f' {endereco}')

        if 'distancia_empresa' in p.text:
            p.text = p.text.replace('distancia_empresa', f' {distancia_empresa}')

        if 'Tempo de deslocamento:' in p.text:
            p.text = p.text.replace('dist_empr_temp', f' {dist_empr_temp}')

        if 'meio_transporte' in p.text:
            p.text = p.text.replace('meio_transporte', f' {meio_transporte}')

        if 'estado_civil' in p.text:
            p.text = p.text.replace('estado_civil', f' {estado_civil}')

        if 'tem_filhos' in p.text:
            p.text = p.text.replace('tem_filhos', f' {tem_filhos}')

        if 'filhos_qtde' in p.text:
            p.text = p.text.replace('filhos_qtde', f' {filhos_qtde}')

        if 'filhos_idade' in p.text:
            p.text = p.text.replace('filhos_idade', f' {filhos_idade}')

        if 'casa_propria' in p.text:
            p.text = p.text.replace('casa_propria', f' {casa_propria}')

        if 'porque_contratado' in p.text:
            p.text = p.text.replace('porque_contratado', f' {porque_contratado}')

        if 'porque_identificou' in p.text:
            p.text = p.text.replace('porque_identificou', f' {porque_identificou}')

        if 'descricao' in p.text:
            p.text = p.text.replace('descricao', f' {descricao}')

        if 'tempo_experiencia' in p.text:
            p.text = p.text.replace('tempo_experiencia', f' {tempo_experiencia}')

        if 'porte_empresa' in p.text:
            p.text = p.text.replace('porte_empresa', f' {porte_empresa}')

        if 'nivel' in p.text:
            p.text = p.text.replace('nivel', f' {nivel}')

        if 'area_atuacao' in p.text:
            p.text = p.text.replace('area_atuacao', f' {area_atuacao}')

        if 'estudando_atualmente' in p.text:
            p.text = p.text.replace('estudando_atualmente', f' {estudando_atualmente}')

        if 'instituicao_curso_periodo' in p.text:
            p.text = p.text.replace('instituicao_curso_periodo', f' {instituicao_curso_periodo}')

        if 'pacote_office' in p.text:
            p.text = p.text.replace('pacote_office', f' {pacote_office}')

        if 'word' in p.text:
            p.text = p.text.replace('word', f' {word}')

        if 'excel' in p.text:
            p.text = p.text.replace('excel', f' {excel}')

        if 'fumante' in p.text:
            p.text = p.text.replace('fumante', f' {fumante}')

        if 'resultado_teste' in p.text:
            p.text = p.text.replace('resultado_teste', f' {resultado_teste}')

        if 'seg_ult_empr' in p.text:
            p.text = p.text.replace('seg_ult_empr', f' {seg_ult_empr}')

        if 'temp_ult_empr' in p.text:
            p.text = p.text.replace('temp_ult_empr', f' {temp_ult_empr}')

        if 'mot_saida_ult_empr' in p.text:
            p.text = p.text.replace('mot_saida_ult_empr', f' {mot_saida_ult_empr}')

        if 'seg_penult_empr' in p.text:
            p.text = p.text.replace('seg_penult_empr', f' {seg_penult_empr}')

        if 'temp_penult_empr' in p.text:
            p.text = p.text.replace('temp_penult_empr', f' {temp_penult_empr}')

        if 'mot_saida_penult_empr' in p.text:
            p.text = p.text.replace('mot_saida_penult_empr', f' {mot_saida_penult_empr}')

        if 'telefone' in p.text:
            p.text = p.text.replace('telefone', f' {telefone}')

        if 'uf_estado' in p.text:
            p.text = p.text.replace('uf_estado', f' {uf_estado}')

        if 'cod_cidade' in p.text:
            p.text = p.text.replace('cod_cidade', f' {cod_cidade}')

    # Localizando a tabela no documento
    table = modelo.tables[0]  # Supondo que a tabela que você deseja preencher seja a primeira do documento

    # Preenchendo as células da tabela com as informações
    table.cell(0, 0).text = "Aguia"
    table.cell(1, 0).text = "Gato"
    table.cell(2, 0).text = "Tubarão"
    table.cell(3, 0).text = "Lobo"
    table.cell(0, 1).text = f"{aguia}%"  # Adiciona o símbolo de porcentagem
    table.cell(1, 1).text = f"{gato}%"
    table.cell(2, 1).text = f"{tubarao}%"
    table.cell(3, 1).text = f"{lobo}%"

    nome_arquivo_sem_acentuacao = unidecode(f'{usuario}_{candidato_vaga}')
#   caminho_arquivo = os.path.join('D:', 'LaudosTeste', f'{nome_arquivo_sem_acentuacao}.docx')

    caminho_arquivo = os.path.join(diretorio_laudos, f'{nome_arquivo_sem_acentuacao}.docx')
    modelo.save(caminho_arquivo)

#    modelo.save(caminho_arquivo)

    print(f'[data: {current_time}].[registro: {carimbo}].[email: {email}].[cod.: ({aguia},{gato},{tubarao},{lobo})].[Laudo {nome_arquivo_sem_acentuacao} criado com sucesso.].[1]')

def calcular_idade(data_nascimento):
    data_nascimento = datetime.strptime(data_nascimento, "%d/%m/%Y")
    data_atual = datetime.now()
    idade = data_atual.year - data_nascimento.year - (
            (data_atual.month, data_atual.day) < (data_nascimento.month, data_nascimento.day))
    return idade

def salvar_estado_laudo(estado):
    with open('estado_laudo.txt', 'w') as arquivo:
        for linha_estado in estado:
            arquivo.write(linha_estado + '\n')

def ler_estado_laudo():
    try:
        with open('estado_laudo.txt', 'r') as arquivo:
            estado = arquivo.read().splitlines()
            return estado
    except FileNotFoundError:
        return []

def gerar_laudos_pendentes():
    planilha = gc.open_by_url(
        "https://docs.google.com/spreadsheets/d/1U9ZFaFmNf4aJ13gbzxuF8AObk-tDWFOaE9If5iOplNc/edit?resourcekey#gid=307731029"). \
        worksheet("Respostas ao formulário 1")
    total_linhas_preenchidas = len(planilha.get_all_values())

    # Obter o estado dos laudos
    laudo_gerado = ler_estado_laudo()

    # Obter os laudos pendentes
    laudos_pendentes = [linha for linha in range(1, total_linhas_preenchidas + 1) if not laudo_ja_gerado(linha)]

    if not laudos_pendentes:
        # Obtém a data e hora atual
        now = datetime.now()
        current_time = now.strftime("%d/%m/%Y %H:%M:%S")
        print(f"[{current_time}] Não há laudos pendentes para gerar.")
    else:
        for linha in laudos_pendentes:
            # Extrair os dados da planilha
            dados = obter_dados_planilha(linha)

            # Preencher o laudo
            carimbo, email, usuario, data_nascimento, candidato_vaga, pretensao_salarial, endereco, distancia_empresa, meio_transporte, estado_civil, tem_filhos, casa_propria, porque_contratado, porque_identificou, descricao, tempo_experiencia, porte_empresa, nivel, area_atuacao, estudando_atualmente, instituicao_curso_periodo, pacote_office, excel, word, fumante, dist_empr_temp, filhos_idade,   = dados[:27]
            candidato_vaga = candidato_vaga.replace('/', 'ou')
            filhos_qtde, seg_ult_empr, temp_ult_empr, mot_saida_ult_empr, seg_penult_empr, temp_penult_empr, mot_saida_penult_empr, telefone, uf_estado, cod_cidade = dados[53:] # BB ate BI
            aguia_percent, gato_percent, tubarao_percent, lobo_percent = contar_respostas(dados)

            # Verificar o animal dominante
            animais = {'I': 'Águia', 'C': 'Gato', 'A': 'Tubarão', 'O': 'Lobo'}
            animal_dominante = max(animais, key=lambda k: dados[27:53].count(k))

            # Normalizar os acentos para suas formas canônicas
            animal_dominante_sem_acento = unicodedata.normalize('NFKD', animais[animal_dominante]).encode('ASCII','ignore').decode('ASCII')

            # Ajustar o resultado do teste de acordo com o animal dominante
            resultado_teste = globals()[f'describe_{animal_dominante_sem_acento.lower()}']

            # Preencher o documento do Word
            preencher_informacoes(carimbo, email, usuario, data_nascimento, candidato_vaga ,pretensao_salarial,
            endereco, distancia_empresa, meio_transporte, estado_civil, tem_filhos, casa_propria, porque_contratado, porque_identificou,
            descricao, tempo_experiencia, porte_empresa, nivel, area_atuacao, estudando_atualmente, instituicao_curso_periodo, pacote_office, excel, word, fumante,
            dist_empr_temp, filhos_idade, filhos_qtde, seg_ult_empr, temp_ult_empr, mot_saida_ult_empr, seg_penult_empr, temp_penult_empr, mot_saida_penult_empr, telefone, uf_estado, cod_cidade,
            aguia_percent, gato_percent, tubarao_percent, lobo_percent, resultado_teste)

            # Gerar o nome do arquivo sem acentuação
            nome_arquivo_sem_acentuacao = unidecode(f'{usuario}_{candidato_vaga}')
            #resultado_teste = unidecode(f'{resultado_teste}')

            # Caminho da imagem que será anexada ao email
            #caminho_imagem = os.path.join('D:', 'LaudosTeste', 'imgLaudos', f'{nome_arquivo_sem_acentuacao}.png')
            caminho_imagem = os.path.join(diretorio_laudos, 'imgLaudos', f'{nome_arquivo_sem_acentuacao}.png')

            # Salvar as informações do candidato para enviar o email posteriormente
            info_candidato = {
                "email": email,
                "nome": usuario,  # Usar a variável correta para nome do candidato
                "candidato_vaga": candidato_vaga,
                "caminho_imagem": caminho_imagem,  # Caminho da imagem do resultado do teste
                "resultado_teste": resultado_teste
            }

            # Salvar essas informações em um arquivo JSON
            with open("candidato_info.json", "w") as file:
                json.dump(info_candidato, file)

            print(f"Informações do candidato {usuario} salvas para envio de email.")

            # Marcar o laudo como gerado
            laudo_gerado.append(str(linha))

        # Salvar o estado dos laudos
        salvar_estado_laudo(laudo_gerado)



def main():
    gerar_laudos_pendentes()

if __name__ == "__main__":
    main()