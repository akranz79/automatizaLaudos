import os
from docx import Document
import logging
from log_config import setup_logging

def inserir_imagem_em_docx(nome_arquivo_docx, caminho_imagem):
    logging.info("Esteira de processo[4.1]: chamando inserir_imagem_em_docx : inserirImgPar42.py")
    # Abre o documento docx
    doc = Document(nome_arquivo_docx)

    # Insere a imagem no parágrafo 42
    paragrafo_index = 42
    if paragrafo_index < len(doc.paragraphs):
        doc.paragraphs[paragrafo_index].add_run().add_picture(caminho_imagem)

        # Salva o documento com a imagem inserida
        doc.save(nome_arquivo_docx)
        #print(f"Imagem do Grafico inserida em {nome_arquivo_docx}.[3]")

    else:
        print(f"O arquivo {nome_arquivo_docx} não possui parágrafo 42.")

def verificar_e_inserir_imagem():
    logging.info("Esteira de processo[4.2]: chamando verificar_e_inserir_imagem : inserirImgPar42.py")
    diretorio_laudos = 'D:/Laudos'
    for arquivo in os.listdir(diretorio_laudos):
        if arquivo.endswith('.docx') and arquivo != 'modelo.docx':
            nome_arquivo_docx = os.path.join(diretorio_laudos, arquivo)
            nome_arquivo_png = os.path.join(diretorio_laudos, arquivo.replace('.docx', '.png'))

            if os.path.exists(nome_arquivo_png):
                inserir_imagem_em_docx(nome_arquivo_docx, nome_arquivo_png)
            else:
                print(f"Não foi encontrado o arquivo PNG correspondente para {nome_arquivo_docx}.")
                logging.warning(f"Não foi encontrado o arquivo PNG correspondente para {nome_arquivo_docx}.")

if __name__ == "__main__":
    verificar_e_inserir_imagem()
